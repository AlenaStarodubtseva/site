from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from io import BytesIO
from docx import Document
import pytz
import datetime
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from flask_migrate import Migrate



app = Flask(__name__)
app.secret_key = "your_secret_key"
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///requests.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)
migrate = Migrate(app, db)

def get_current_time_japan():
    japan_timezone = pytz.timezone('Asia/Tokyo')
    current_time = datetime.datetime.now(japan_timezone)
    return current_time  # Вернем объект datetime, а не строку

# Модель данных
class Request(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    course_group = db.Column(db.String(20), nullable=False)
    destination = db.Column(db.String(100), nullable=False)
    request_type = db.Column(db.String(50), nullable=False)
    quantity = db.Column(db.Integer, nullable=False)
    period_start = db.Column(db.Date, nullable=True)
    period_end = db.Column(db.Date, nullable=True)
    created_date = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class ArchiveWithScholarship(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    request_id = db.Column(db.Integer, nullable=False)
    processed_date = db.Column(db.DateTime, nullable=False, default=datetime.datetime.utcnow)
    admin_comment = db.Column(db.String(200), nullable=True)
    name = db.Column(db.String(200), nullable=False)  # Это поле отсутствует в таблице
    course_group = db.Column(db.String(10), nullable=False)
    destination = db.Column(db.String(200), nullable=False)
    quantity = db.Column(db.Integer, nullable=False)
    period_start = db.Column(db.Date, nullable=True)
    period_end = db.Column(db.Date, nullable=True)
    request_type = db.Column(db.String(100), nullable=False)


class ArchiveWithoutScholarship(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    request_id = db.Column(db.Integer, nullable=False)  # ID из основной таблицы
    processed_date = db.Column(db.DateTime, default=datetime.datetime.utcnow)  # Дата обработки
    admin_comment = db.Column(db.String(200), nullable=True)  # Комментарий администратора
    name = db.Column(db.String(200), nullable=False)  # Имя заявителя
    course_group = db.Column(db.String(50), nullable=False)  # Курс и группа
    destination = db.Column(db.String(200), nullable=False)  # Место, куда нужна справка
    quantity = db.Column(db.Integer, nullable=False)  # Количество справок
    period_start = db.Column(db.Date, nullable=True)  # Период начала (если применимо)
    period_end = db.Column(db.Date, nullable=True)  # Период окончания (если применимо)
    request_type = db.Column(db.String(100), nullable=False)  # Тип справки


with app.app_context():
    db.create_all()


# Главная страница (заявка)
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        name = request.form.get("name")
        course = int(request.form["course"])
        group = request.form.get("group")
        destination = request.form.get("destination")
        request_type = request.form.get("request_type")
        quantity = int(request.form["quantity"])
        period_start = request.form.get("period_start")
        period_end = request.form.get("period_end")

            
        # Проверка на тип справки
        if request_type == "Справка с отметкой о стипендии" and (not period_start or not period_end):
            flash("Пожалуйста, укажите период для справки с отметкой о стипендии.")
            return redirect("/")

        # Проверка курсов для групп
        if group in ["МФМО", "МХПО"] and course > 2:
            flash("Группы МФМО и МХПО не могут быть на курсе больше 2.", "error")
            return redirect(url_for("index"))

        if group in ["МО", "ИС", "ЭУ", "СПД", "СПИ"] and course > 4:
            flash("Группы МО, ИС, ЭУ, СПД, СПИ не могут быть на курсе больше 4.", "error")
            return redirect(url_for("index"))

        # Проверка периода
        if period_start and period_end and period_start > period_end:
            flash("Начальная дата не может быть позже конечной.", "error")
            return redirect(url_for("index"))

        # Проверка количества справок
        if quantity < 1:
            flash("Количество справок должно быть не меньше 1.", "error")
            return redirect(url_for("index"))
        
        # Создание новой заявки
        new_request = Request(
            name=name,
            course_group=f"{course}{group}",
            destination=destination,
            request_type=request_type,
            quantity=quantity,
            period_start=datetime.datetime.strptime(period_start, "%Y-%m-%d") if period_start else None,
            period_end=datetime.datetime.strptime(period_end, "%Y-%m-%d") if period_end else None,
        )
        db.session.add(new_request)
        db.session.commit()
        flash("Заявка успешно отправлена!")
        return redirect("/")
    return render_template("index.html")


# Админ-панель
@app.route("/admin")
def admin():
    requests_with_stipend = Request.query.filter_by(request_type="Справка с отметкой о стипендии").all()
    requests_without_stipend = Request.query.filter_by(request_type="Справка без отметки о стипендии").all()
    return render_template("admin.html", requests_with_stipend=requests_with_stipend, requests_without_stipend=requests_without_stipend)


@app.route('/archive/without-scholarship', methods=['GET', 'POST'])
def archive_without_scholarship():
    if request.method == 'POST':
        # Получение выбранных заявок из архива
        selected_ids = request.form.getlist("selected_ids")
        action = request.form.get("action")
        if action == "return":
            # Возврат заявок в админ-панель
            archived_requests = ArchiveWithoutScholarship.query.filter(ArchiveWithoutScholarship.id.in_(selected_ids)).all()
            for archived_request in archived_requests:
                # Перенос из архива в Requests
                original_request = Request(
                    id=archived_request.request_id,
                    name=archived_request.name,
                    course_group=archived_request.course_group,
                    destination=archived_request.destination,
                    request_type="Справка без отметки о стипендии",
                    quantity=archived_request.quantity,
                    period_start=archived_request.period_start,
                    period_end=archived_request.period_end,
                )
                db.session.add(original_request)
                db.session.delete(archived_request)
            db.session.commit()
            flash("Заявки успешно возвращены в админ-панель", "success")
        elif action == "delete":
            # Удаление выбранных заявок
            ArchiveWithoutScholarship.query.filter(ArchiveWithoutScholarship.id.in_(selected_ids)).delete()
            db.session.commit()
            flash("Заявки успешно удалены", "success")
    archived_requests = ArchiveWithoutScholarship.query.all()
    return render_template("archive_without_scholarship.html", requests=archived_requests)


@app.route('/archive/with-scholarship', methods=['GET', 'POST'])
def archive_with_scholarship():
    if request.method == 'POST':
        # Получение выбранных заявок из архива
        selected_ids = request.form.getlist("selected_ids")
        action = request.form.get("action")
        if action == "return":
            # Возврат заявок в админ-панель
            archived_requests = ArchiveWithScholarship.query.filter(ArchiveWithScholarship.id.in_(selected_ids)).all()
            for archived_request in archived_requests:
                # Перенос из архива в Requests
                original_request = Request(
                    id=archived_request.request_id,
                    name=archived_request.name,
                    course_group=archived_request.course_group,
                    destination=archived_request.destination,
                    request_type="Справка с отметкой о стипендии",
                    quantity=archived_request.quantity,
                    period_start=archived_request.period_start,
                    period_end=archived_request.period_end,
                )
                db.session.add(original_request)
                db.session.delete(archived_request)
            db.session.commit()
            flash("Заявки успешно возвращены в админ-панель", "success")
        elif action == "delete":
            # Удаление выбранных заявок
            ArchiveWithScholarship.query.filter(ArchiveWithScholarship.id.in_(selected_ids)).delete()
            db.session.commit()
            flash("Заявки успешно удалены", "success")
    archived_requests = ArchiveWithScholarship.query.all()
    return render_template("archive_with_scholarship.html", requests=archived_requests)


# без отметки
@app.route("/delete", methods=["POST"])
def handle_requests_without_stipend():
    selected_ids = request.form.getlist("delete_ids")
    action = request.form.get("action")

    if action == "delete":
        # Удаление заявок
        requests_to_delete = Request.query.filter(Request.id.in_(selected_ids)).all()
        for req in requests_to_delete:
            db.session.delete(req)
        db.session.commit()
        flash("Выбранные заявки удалены.", "success")
    elif action == "archive":
        # Архивирование заявок
        requests_to_archive = Request.query.filter(Request.id.in_(selected_ids)).all()
        for req in requests_to_archive:
            archived_request = ArchiveWithoutScholarship(
                request_id=req.id,
                processed_date=datetime.datetime.now(),
                admin_comment="Справка добавлена в архив",
                name=req.name,
                course_group=req.course_group,
                destination=req.destination,
                quantity=req.quantity,
                period_start=req.period_start,
                period_end=req.period_end,
                request_type=req.request_type
            )
            db.session.add(archived_request)
            db.session.delete(req)  # Удаляем из основной таблицы
            
        db.session.commit()  # Сохраняем изменения
        flash("Выбранные заявки архивированы.", "success")
        return redirect(url_for('admin'))
    elif action == "view_archive":
        # Перенаправление на архив
        return redirect("/archive/without-scholarship")

    return redirect("/admin")

# с отметкой
@app.route('/handle-requests', methods=['POST'])
def handle_requests():
    selected_ids = request.form.getlist('request_ids')  # Получаем выбранные ID заявок
    action = request.form.get('action')  # Получаем значение action

    if action == "delete":
        # Удаление заявок из базы данных
        for request_id in selected_ids:
            request_to_delete = Request.query.get(request_id)
            if request_to_delete:
                db.session.delete(request_to_delete)
        db.session.commit()
        flash("Выбранные заявки успешно удалены", "success")
        return redirect(url_for('admin'))

    elif action == "archive":
        # Архивирование заявок
        requests_to_archive = Request.query.filter(Request.id.in_(selected_ids)).all()
        for req in requests_to_archive:
            archived_request = ArchiveWithScholarship(
                request_id=req.id,
                processed_date=datetime.datetime.now(),
                admin_comment="Справка добавлена в архив",
                name=req.name,
                course_group=req.course_group,
                destination=req.destination,
                quantity=req.quantity,
                period_start=req.period_start,
                period_end=req.period_end,
                request_type=req.request_type
            )
            db.session.add(archived_request)
            db.session.delete(req)  # Удаляем из основной таблицы
            
        db.session.commit()  # Сохраняем изменения
        return redirect(url_for('admin'))
    elif action == "view_archive":
        # Перенаправление на архив
        return redirect("/archive/with-scholarship")
    elif action == "generate":
        # Генерация общего документа
        selected_requests = Request.query.filter(Request.id.in_(selected_ids)).all()
        doc = Document()
        
        # Установка шрифта для всего документа
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Заголовок документа
        title = doc.add_heading(level=1)
        run = title.add_run(f"Заявка на справки {get_current_time_japan().strftime('%d.%m.%Y')} факультета физико-математического образования и технологии")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Создание таблицы
        table = doc.add_table(rows=1, cols=5)  # Заголовки таблицы
        table.style = 'Table Grid'
        
        # Добавление заголовков
        hdr_cells = table.rows[0].cells
        hdr_titles = ["Ф.И.О.", "Курс, группа", "Куда", "Сколько", "Примечание"]
        for idx, cell in enumerate(hdr_cells):
            cell.text = hdr_titles[idx]
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

        # Установка стиля заголовков таблицы
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # Заполнение таблицы данными
        for req in selected_requests:
            row_cells = table.add_row().cells
            row_cells[0].text = req.name
            row_cells[1].text = req.course_group
            row_cells[2].text = req.destination
            row_cells[3].text = str(req.quantity)

        # Добавление примечания для справки с отметкой о стипендии
            if req.request_type == "Справка с отметкой о стипендии":
                row_cells[4].text = f"{req.period_start.strftime('%d.%m.%Y')} – {req.period_end.strftime('%d.%m.%Y')}"
            else:
                row_cells[4].text = "-"

        # Форматирование текста в каждой ячейке строки
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
        # Сохранение документа во временный файл
        temp_file = BytesIO()
        doc.save(temp_file)
        temp_file.seek(0)

        return send_file(
            temp_file,
            as_attachment=True,
            download_name="Общий_документ.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    else:
        flash("Неизвестное действие", "error")
        return redirect(url_for('admin'))


